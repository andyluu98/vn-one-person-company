"""Render document tu template + data -> .docx hoac .xlsx."""
from __future__ import annotations
from pathlib import Path
from typing import Any


class DocWriter:
    def __init__(self, output_root: Path):
        self.output_root = Path(output_root)

    def write_docx(
        self,
        template_path: Path,
        output_rel: str,
        substitutions: dict[str, Any],
    ) -> Path:
        from docx import Document

        out_path = self.output_root / output_rel
        out_path.parent.mkdir(parents=True, exist_ok=True)

        if template_path.suffix == ".docx":
            doc = Document(str(template_path))
            self._substitute_docx(doc, substitutions)
        else:
            doc = Document()
            text = template_path.read_text(encoding="utf-8")
            text = self._substitute(text, substitutions)
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)

        doc.save(str(out_path))
        return out_path

    def write_xlsx(
        self,
        template_path: Path,
        output_rel: str,
        rows: list[dict],
    ) -> Path:
        from openpyxl import load_workbook, Workbook

        out_path = self.output_root / output_rel
        out_path.parent.mkdir(parents=True, exist_ok=True)

        if template_path.suffix == ".xlsx" and template_path.exists():
            wb = load_workbook(str(template_path))
        else:
            wb = Workbook()

        ws = wb.active
        if rows and ws.max_row == 1 and ws.cell(1, 1).value is None:
            headers = list(rows[0].keys())
            for col, h in enumerate(headers, 1):
                ws.cell(1, col, h)

        if rows:
            headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
            for row in rows:
                next_row = ws.max_row + 1
                for col, h in enumerate(headers, 1):
                    ws.cell(next_row, col, row.get(h, ""))

        wb.save(str(out_path))
        return out_path

    @staticmethod
    def _substitute(text: str, subs: dict[str, Any]) -> str:
        for k, v in subs.items():
            text = text.replace(f"{{{{{k}}}}}", str(v))
        return text

    @staticmethod
    def _substitute_docx(doc, subs: dict):
        for para in doc.paragraphs:
            for k, v in subs.items():
                if f"{{{{{k}}}}}" in para.text:
                    for run in para.runs:
                        run.text = run.text.replace(f"{{{{{k}}}}}", str(v))
